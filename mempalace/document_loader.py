"""Read project source files into plain text for indexing.

This module is intentionally read-only: it never mutates source files.
"""

from __future__ import annotations

from pathlib import Path

from openpyxl import load_workbook
from pypdf import PdfReader
from docx import Document

from .asset_store import IMAGE_EXTENSIONS


TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".py",
    ".js",
    ".ts",
    ".jsx",
    ".tsx",
    ".json",
    ".yaml",
    ".yml",
    ".html",
    ".css",
    ".java",
    ".go",
    ".rs",
    ".rb",
    ".sh",
    ".csv",
    ".sql",
    ".toml",
}

EXTRACTED_DOCUMENT_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".xlsx",
}

METADATA_ONLY_EXTENSIONS = set(IMAGE_EXTENSIONS)

SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | EXTRACTED_DOCUMENT_EXTENSIONS | METADATA_ONLY_EXTENSIONS


def read_source_text(filepath: Path) -> str:
    """Return text content for a file, extracting from rich docs when needed."""
    suffix = filepath.suffix.lower()

    if suffix in TEXT_EXTENSIONS:
        return filepath.read_text(encoding="utf-8", errors="replace")
    if suffix == ".pdf":
        return _read_pdf(filepath)
    if suffix == ".docx":
        return _read_docx(filepath)
    if suffix == ".xlsx":
        return _read_xlsx(filepath)
    if suffix in METADATA_ONLY_EXTENSIONS:
        return _binary_placeholder(filepath, note="Image asset; OCR not enabled.")

    raise ValueError(f"Unsupported file type: {suffix}")


def _read_pdf(filepath: Path) -> str:
    reader = PdfReader(str(filepath))
    sections = [f"[PDF] {filepath.name}"]

    for page_number, page in enumerate(reader.pages, 1):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        text = _clean_text(text)
        if text:
            sections.append(f"[Page {page_number}]")
            sections.append(text)

    if len(sections) == 1:
        return _binary_placeholder(filepath, note="PDF had no extractable text.")
    return "\n\n".join(sections)


def _read_docx(filepath: Path) -> str:
    document = Document(str(filepath))
    lines = [f"[DOCX] {filepath.name}"]

    for paragraph in document.paragraphs:
        text = _clean_text(paragraph.text)
        if text:
            lines.append(text)

    for table_index, table in enumerate(document.tables, 1):
        lines.append(f"[Table {table_index}]")
        for row in table.rows:
            cells = [_clean_text(cell.text) for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                lines.append(" | ".join(cells))

    if len(lines) == 1:
        return _binary_placeholder(filepath, note="DOCX had no extractable text.")
    return "\n\n".join(lines)


def _read_xlsx(filepath: Path) -> str:
    workbook = load_workbook(filename=str(filepath), read_only=True, data_only=True)
    lines = [f"[XLSX] {filepath.name}"]

    for sheet in workbook.worksheets:
        lines.append(f"[Sheet] {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = []
            for value in row:
                if value is None:
                    continue
                text = _clean_text(str(value))
                if text:
                    cells.append(text)
            if cells:
                lines.append(" | ".join(cells))

    workbook.close()

    if len(lines) <= 2:
        return _binary_placeholder(filepath, note="Workbook had no extractable cell text.")
    return "\n\n".join(lines)


def _binary_placeholder(filepath: Path, note: str) -> str:
    stat = filepath.stat()
    return "\n".join(
        [
            f"[FILE] {filepath.name}",
            f"Type: {filepath.suffix.lower()}",
            f"Path: {filepath}",
            f"SizeBytes: {stat.st_size}",
            f"Note: {note}",
        ]
    )


def _clean_text(text: str) -> str:
    stripped_lines = [line.strip() for line in text.replace("\x00", " ").splitlines()]
    kept = [line for line in stripped_lines if line]
    return "\n".join(kept)
