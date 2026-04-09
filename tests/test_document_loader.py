from pathlib import Path

from docx import Document
from openpyxl import Workbook
from pypdf import PdfWriter

from mempalace.document_loader import SUPPORTED_EXTENSIONS, read_source_text


def test_read_docx_extracts_text_and_tables(tmp_path):
    file_path = tmp_path / "notes.docx"
    document = Document()
    document.add_paragraph("Paragraph text")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "A"
    table.rows[0].cells[1].text = "B"
    document.save(file_path)

    text = read_source_text(file_path)

    assert "[DOCX] notes.docx" in text
    assert "Paragraph text" in text
    assert "A | B" in text


def test_read_xlsx_extracts_cells(tmp_path):
    file_path = tmp_path / "sheet.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Data"
    sheet["A1"] = "Name"
    sheet["B1"] = "Value"
    sheet["A2"] = "foo"
    sheet["B2"] = 42
    workbook.save(file_path)

    text = read_source_text(file_path)

    assert "[XLSX] sheet.xlsx" in text
    assert "[Sheet] Data" in text
    assert "Name | Value" in text
    assert "foo | 42" in text


def test_read_pdf_without_extractable_text_returns_placeholder(tmp_path):
    file_path = tmp_path / "blank.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    with open(file_path, "wb") as f:
        writer.write(f)

    text = read_source_text(file_path)

    assert "[FILE] blank.pdf" in text
    assert "PDF had no extractable text." in text


def test_image_extensions_are_supported_for_metadata_only_indexing(tmp_path):
    file_path = tmp_path / "figure.gif"
    file_path.write_bytes(b"GIF89a")

    text = read_source_text(file_path)

    assert ".gif" in SUPPORTED_EXTENSIONS
    assert "figure.gif" in text
    assert "OCR not enabled" in text
